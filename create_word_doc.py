#!/usr/bin/env python3
"""
Create SUMA GUI Documentation Word Document
Script to create a comprehensive Word document with screenshots
"""

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.shared import RGBColor
    import os
    
    def create_suma_documentation():
        """Create comprehensive SUMA documentation"""
        print("Creating SUMA GUI Documentation...")
        
        # Create new document
        doc = Document()
        
        # Add title
        title = doc.add_heading('SUMA GUI - תיעוד מקיף', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_heading('Signal Utility MRI Analysis - Professional GUI', level=1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add overview section
        doc.add_heading('סקירה כללית', level=1)
        
        doc.add_paragraph(
            'SUMA (Signal Utility MRI Analysis) היא תוכנת שולחן עבודה מקצועית המיועדת '
            'לניתוח תמונות MRI ו-CT של מטופלים במחקרים קליניים. התוכנה מחשבת באופן אוטומטי:'
        )
        
        overview_list = doc.add_paragraph()
        overview_list.add_run('• TE₀ - זמן הדהוי הטבעי\n')
        overview_list.add_run('• צבירת ברזל ברקמות\n')
        overview_list.add_run('• שינויים לאורך זמן בין טיפולים')
        
        # Add workflow section
        doc.add_heading('זרימת העבודה', level=1)
        
        doc.add_paragraph(
            'התוכנה פועלת על פי זרימת עבודה פשוטה ואינטואיטיבית:'
        )
        
        workflow = doc.add_paragraph()
        workflow.add_run('1. התחלת הניתוח: מסך פתיחה → דיאלוג ניתוח חדש → הזנת פרטי מטופל → בחירת תיקיות → ביצוע ניתוח\n')
        workflow.add_run('2. צפייה בתוצאות: מסך פתיחה → דיאלוג צפייה בתוצאות → בחירת מטופל/טיפול → מציג תוצאות\n')
        workflow.add_run('3. ייצוא תוצאות: מציג תוצאות → כפתור Export PDF → יצירת דוח מקצועי')
        
        # Add screenshots section
        doc.add_heading('מסכי המערכת', level=1)
        
        # Screenshot 1 - Startup Screen
        doc.add_heading('1. מסך הפתיחה (Startup Screen)', level=2)
        doc.add_paragraph(
            'המסך הראשי של התוכנה המציג את לוגו SUMA ושני כפתורי פעולה עיקריים.'
        )
        
        if os.path.exists('screenshots/01_startup_screen.png'):
            doc.add_picture('screenshots/01_startup_screen.png', width=Inches(6))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        startup_features = doc.add_paragraph()
        startup_features.add_run('רכיבים עיקריים:\n')
        startup_features.add_run('• כותרת ראשית "SUMA" בגופן גדול (96px) בצבע כחול כהה\n')
        startup_features.add_run('• כותרת משנה "Wellcome to Signal Utility MRI Analysis" בגופן בינוני (48px)\n')
        startup_features.add_run('• כפתור "Start Analysis" להתחלת ניתוח חדש\n')
        startup_features.add_run('• כפתור "View Results" לצפייה בתוצאות קיימות')
        
        # Screenshot 2 - Analysis Dialog
        doc.add_heading('2. דיאלוג התחלת ניתוח (Start Analysis Dialog)', level=2)
        doc.add_paragraph(
            'דיאלוג מקיף לביצוע ניתוח חדש על נתוני מטופל.'
        )
        
        if os.path.exists('screenshots/02_analysis_dialog.png'):
            doc.add_picture('screenshots/02_analysis_dialog.png', width=Inches(5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        analysis_features = doc.add_paragraph()
        analysis_features.add_run('פונקציונליות מרכזית:\n')
        analysis_features.add_run('• הזנת פרטי מטופל (Patient ID, Treatment ID)\n')
        analysis_features.add_run('• בחירת תיקיות נתונים (MRI PRE, MRI POST, CT)\n')
        analysis_features.add_run('• Progress Bar עם התקדמות באחוזים\n')
        analysis_features.add_run('• השבתת רכיבים במהלך הניתוח\n')
        analysis_features.add_run('• ביצוע ברקע עם Worker Thread')
        
        # Screenshot 3 - Results Dialog
        doc.add_heading('3. דיאלוג צפייה בתוצאות (View Results Dialog)', level=2)
        doc.add_paragraph(
            'דיאלוג המציג רשימת מטופלים וטיפולים קיימים במערכת.'
        )
        
        if os.path.exists('screenshots/03_results_dialog.png'):
            doc.add_picture('screenshots/03_results_dialog.png', width=Inches(5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        results_features = doc.add_paragraph()
        results_features.add_run('מאפיינים:\n')
        results_features.add_run('• תצוגה היררכית של מטופלים וטיפולים\n')
        results_features.add_run('• יישור שמאל לקריאות טובה יותר\n')
        results_features.add_run('• בחירה אינטראקטיבית עם לחיצה כפולה\n')
        results_features.add_run('• מידע מפורט: תאריך וסטטוס לכל טיפול')
        
        # Screenshot 4 - Results Viewer
        doc.add_heading('4. מציג תוצאות (Results Viewer)', level=2)
        doc.add_paragraph(
            'המסך הראשי לצפייה ועריכת תוצאות הניתוח.'
        )
        
        if os.path.exists('screenshots/04_results_viewer.png'):
            doc.add_picture('screenshots/04_results_viewer.png', width=Inches(6))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        viewer_features = doc.add_paragraph()
        viewer_features.add_run('רכיבים מתקדמים:\n')
        viewer_features.add_run('• סרגל עליון עם מידע מטופל וכפתורי פעולה\n')
        viewer_features.add_run('• אזור תצוגה מוכן לאינטגרציה עם Napari\n')
        viewer_features.add_run('• שכבות תצוגה: CT, PRE TE₁, POST TE₁, PRE TE₀, POST TE₀, ΔTE₀, Iron\n')
        viewer_features.add_run('• בקרים: Slice slider, 3D toggle, Export PDF')
        
        # Add technical specifications
        doc.add_heading('מפרט טכני', level=1)
        
        tech_specs = doc.add_paragraph()
        tech_specs.add_run('דרישות מערכת:\n')
        tech_specs.add_run('• מערכת הפעלה: Windows 10/11\n')
        tech_specs.add_run('• Python: 3.8+\n')
        tech_specs.add_run('• זיכרון: 8GB RAM מינימום\n')
        tech_specs.add_run('• מקום פנוי: 2GB לתוכנה ונתונים\n\n')
        
        tech_specs.add_run('תלות טכנית:\n')
        tech_specs.add_run('• PyQt5 - ממשק משתמש\n')
        tech_specs.add_run('• NumPy - עיבוד מטריצות\n')
        tech_specs.add_run('• Napari - תצוגת תמונות (אינטגרציה עתידית)')
        
        # Add color scheme
        doc.add_heading('ערכת צבעים', level=1)
        
        colors = doc.add_paragraph()
        colors.add_run('התוכנה משתמשת בפלטת צבעים מוגדרת:\n')
        colors.add_run('• כחול כהה (#004466) - צבע מותג ראשי\n')
        colors.add_run('• כחול בהיר (#007acc) - הדגשות וכפתורים\n')
        colors.add_run('• לבן (#FFFFFF) - רקעים\n')
        colors.add_run('• אפור כהה (#333333) - טקסט ראשי\n')
        colors.add_run('• אפור בהיר (#999999) - טקסט משני')
        
        # Add summary
        doc.add_heading('סיכום', level=1)
        
        summary = doc.add_paragraph(
            'SUMA GUI היא תוכנה מקצועית ומתקדמת לניתוח תמונות רפואיות במחקרים קליניים. '
            'התוכנה מציעה ממשק אינטואיטיבי, ניתוח אוטומטי, תצוגה מתקדמת, ייצוא מקצועי '
            'וארכיטקטורה גמישה. התוכנה מהווה כלי חיוני למחקרים קליניים העוסקים בצבירת ברזל '
            'ובהשפעות טיפולים על רקמות המוח.'
        )
        
        # Add footer
        footer = doc.add_paragraph('© 2025 New Phase Ltd. - For internal research use only')
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Save document
        doc.save('SUMA_GUI_Documentation.docx')
        print("✅ Document saved: SUMA_GUI_Documentation.docx")
        
    if __name__ == "__main__":
        create_suma_documentation()
        
except ImportError:
    print("❌ python-docx not installed. Installing...")
    import subprocess
    import sys
    
    try:
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
        print("✅ python-docx installed successfully!")
        print("Please run the script again.")
    except subprocess.CalledProcessError:
        print("❌ Failed to install python-docx")
        print("Please install manually: pip install python-docx") 